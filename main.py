#!/usr/bin/env python3
"""
Audio Transcription & Documentation Tool

Processes audio files by compressing them, transcribing with OpenAI Whisper,
formatting with GPT-4, and saving as Word documents (.docx).
"""

import os
import json
import argparse
import logging
import traceback
import time
from pathlib import Path
from typing import Optional, Dict, Tuple
from datetime import datetime

import httpx
import ffmpeg
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('transcription_errors.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Constants
INPUT_DIR = "input"
COMPRESSED_DIR = "compressed"
OUTPUT_DIR = "output"
SUPPORTED_FORMATS = ['.mp3', '.wav', '.m4a', '.flac', '.aac']
MAX_FILE_SIZE_MB = 25
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

# Initialize OpenAI client with granular timeout settings for large file uploads
# Different phases of HTTP requests need different timeouts:
# - connect: Time to establish connection (short)
# - read: Time to receive response (long for Whisper processing)
# - write: Time to upload file (long for large audio files)
# - pool: Time to acquire connection from pool (short)
timeout_config = httpx.Timeout(
    connect=30.0,    # 30 seconds to establish connection
    read=900.0,      # 15 minutes to receive response (Whisper processing)
    write=300.0,     # 5 minutes to upload file
    pool=30.0        # 30 seconds to acquire connection from pool
)

client = OpenAI(
    api_key=os.getenv('OPENAI_API_KEY'),
    timeout=timeout_config,
    max_retries=0   # Disable SDK retries, use application-level retry logic only
)


def check_network_connectivity() -> bool:
    """
    Check if we can reach OpenAI API endpoints.

    Returns:
        True if network is reachable, False otherwise
    """
    import socket
    try:
        # Try to resolve DNS for api.openai.com
        socket.gethostbyname('api.openai.com')
        logger.info("Network connectivity check: OK")
        return True
    except socket.error as e:
        logger.error(f"Network connectivity check FAILED: {e}")
        print(f"WARNING: Cannot reach api.openai.com - check your internet connection")
        return False


def compress_audio(input_path: str, output_path: str) -> bool:
    """
    Compress audio file to MP3 using FFmpeg with adaptive bitrate.
    Tries progressively lower bitrates until file is under 25MB.

    Args:
        input_path: Path to input audio file
        output_path: Path to save compressed audio

    Returns:
        True if successful, False otherwise
    """
    # Try progressively lower bitrates until file fits under 25MB
    bitrates = ['64k', '48k', '32k', '24k']

    for bitrate in bitrates:
        try:
            if bitrate == '64k':
                print(f"  Compressing audio...")
            else:
                print(f"  Re-compressing at {bitrate} bitrate...")

            (
                ffmpeg
                .input(input_path)
                .output(
                    output_path,
                    ar=16000,      # Sample rate: 16kHz
                    ac=1,          # Channels: Mono
                    **{'b:a': bitrate}
                )
                .overwrite_output()
                .run(capture_stdout=True, capture_stderr=True, quiet=True)
            )

            # Check if file is under the limit
            size_bytes = os.path.getsize(output_path)
            size_mb = size_bytes / (1024 * 1024)

            if size_mb < MAX_FILE_SIZE_MB:
                return True
            else:
                print(f"  Compressed size: {size_mb:.2f}MB (still too large)")
                # Continue to try lower bitrate

        except ffmpeg.Error as e:
            print(f"  ERROR: FFmpeg compression failed: {e.stderr.decode()}")
            return False
        except Exception as e:
            print(f"  ERROR: Compression failed: {str(e)}")
            return False

    # If we've tried all bitrates and still too large
    print(f"  ERROR: File still too large after all compression attempts")
    return False


def check_file_size(file_path: str) -> Tuple[bool, float]:
    """
    Check if file is under the size limit for Whisper API.

    Args:
        file_path: Path to file to check

    Returns:
        Tuple of (is_valid, size_in_mb)
    """
    size_bytes = os.path.getsize(file_path)
    size_mb = size_bytes / (1024 * 1024)
    is_valid = size_bytes < MAX_FILE_SIZE_BYTES
    return is_valid, size_mb


def transcribe_audio(audio_path: str, max_retries: int = 5) -> Optional[str]:
    """
    Transcribe audio file using OpenAI Whisper API with robust retry logic.

    Args:
        audio_path: Path to audio file
        max_retries: Maximum number of retry attempts (default 5)

    Returns:
        Transcription text or None if failed
    """
    for attempt in range(max_retries):
        try:
            if attempt > 0:
                print(f"  Retry attempt {attempt + 1}/{max_retries}...")
            else:
                print(f"  Transcribing audio (this may take 2-5 minutes for long files)...")

            start_time = time.time()

            with open(audio_path, 'rb') as audio_file:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    response_format="text"
                )

            elapsed = time.time() - start_time
            print(f"  Transcription completed in {elapsed:.1f} seconds")
            return transcript

        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__

            # Log detailed error information
            logger.error(f"Transcription attempt {attempt + 1}/{max_retries} failed")
            logger.error(f"Error type: {error_type}")
            logger.error(f"Error message: {error_msg}")
            logger.error(f"File: {audio_path}")

            # Get file size for diagnostics
            try:
                file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
                logger.error(f"File size: {file_size_mb:.2f} MB")
            except:
                pass

            # Check if it's a rate limit error
            if 'rate_limit' in error_msg.lower():
                wait_time = 60 * (attempt + 1)  # 60, 120, 180 seconds for rate limits
                print(f"  Rate limit hit. Waiting {wait_time} seconds...")
                logger.warning(f"Rate limit encountered, waiting {wait_time}s")
            elif '502' in error_msg or 'bad gateway' in error_msg.lower():
                wait_time = min(60 * (attempt + 1), 300)  # 60, 120, 180, 240, 300s for server errors
                print(f"  OpenAI server error (502 Bad Gateway). Waiting {wait_time} seconds...")
                logger.warning(f"502 Bad Gateway from OpenAI, waiting {wait_time}s")
            elif 'timeout' in error_msg.lower() or 'timed out' in error_msg.lower():
                wait_time = min(30 * (2 ** attempt), 180)
                print(f"  Timeout error. Waiting {wait_time} seconds...")
                logger.warning(f"Timeout error, waiting {wait_time}s")
            elif 'connection' in error_msg.lower():
                wait_time = min(20 * (2 ** attempt), 180)
                print(f"  Connection error. Waiting {wait_time} seconds...")
                logger.warning(f"Connection error, waiting {wait_time}s")
            else:
                # Exponential backoff: 10, 20, 40, 80, 160 seconds
                wait_time = min(10 * (2 ** attempt), 180)  # Cap at 3 minutes
                print(f"  Error ({error_type}): {error_msg}")

            if attempt < max_retries - 1:
                print(f"  Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                print(f"  ERROR: Transcription failed after {max_retries} attempts: {error_msg}")
                logger.error(f"Transcription FAILED after all retries: {audio_path}")
                return None


def add_paragraph_breaks(text: str) -> str:
    """
    Add basic paragraph breaks to transcript (fallback for very long transcripts).

    Args:
        text: Raw transcript text

    Returns:
        Text with paragraph breaks added
    """
    import re

    # Split on sentence endings
    sentences = re.split(r'([.!?]+\s+)', text)

    paragraphs = []
    current_para = []
    sentence_count = 0

    for i in range(0, len(sentences), 2):
        sentence = sentences[i] if i < len(sentences) else ''
        punct = sentences[i+1] if i+1 < len(sentences) else ''

        if sentence.strip():
            current_para.append(sentence + punct)
            sentence_count += 1

        # New paragraph every 5 sentences
        if sentence_count >= 5:
            paragraphs.append(''.join(current_para).strip())
            current_para = []
            sentence_count = 0

    if current_para:
        paragraphs.append(''.join(current_para).strip())

    return '\n\n'.join(paragraphs)


def format_with_gpt4(transcript: str, max_retries: int = 3) -> Optional[Dict[str, str]]:
    """
    Format transcript using GPT-5-mini to generate title, summary, and formatted transcript.
    Uses gpt-5-mini with 128K output tokens (can handle ~6-8 hours of audio) at 80% lower cost than gpt-5.
    Falls back to basic formatting if transcript is too long.

    Args:
        transcript: Raw transcript text
        max_retries: Maximum number of retry attempts (default 3)

    Returns:
        Dictionary with 'title', 'summary', and 'transcript' keys, or None if failed
    """
    for attempt in range(max_retries):
        try:
            if attempt > 0:
                print(f"  Retry attempt {attempt + 1}/{max_retries} for formatting...")
            else:
                print(f"  Formatting with GPT-5-mini...")

            start_time = time.time()

            prompt = f"""You are a professional editor. Format the following transcript into a structured document.

YOUR TASK:
1. Keep EVERY SINGLE WORD from the transcript - do not summarize, shorten, or omit ANY content
2. Organize the words into logical paragraphs by adding \\n\\n between paragraph breaks
3. Group related sentences together - each major topic or idea gets its own paragraph

IMPORTANT: Respond ONLY with valid JSON. Do not include markdown code blocks or explanations.

Required JSON structure:
{{
  "title": "A descriptive 5-10 word title for this content",
  "summary": "A concise 2-3 sentence summary of the main points",
  "transcript": "The full transcript with all words preserved. Separate paragraphs with \\n\\n. Make logical paragraph breaks based on topic changes and natural flow."
}}

Raw transcript:
{transcript}

Return ONLY the JSON object. The transcript must contain ALL the original words, just organized into paragraphs with \\n\\n separators."""

            response = client.chat.completions.create(
                model="gpt-5-mini",  # Using gpt-5-mini for massive output token limit (128,000 tokens) at lower cost
                messages=[
                    {"role": "system", "content": "You are a professional editor that outputs only valid JSON. You NEVER shorten or summarize transcripts - you preserve every word."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                max_completion_tokens=128000  # Maximum output tokens for gpt-5-mini (can handle ~6-8 hours of audio)
            )

            content = response.choices[0].message.content
            if not content:
                raise ValueError("Empty response from GPT-5")

            result = json.loads(content)

            elapsed = time.time() - start_time
            print(f"  Formatting completed in {elapsed:.1f} seconds")

            # Validate required keys
            if not all(key in result for key in ['title', 'summary', 'transcript']):
                print(f"  ERROR: GPT-5 response missing required keys")
                if attempt < max_retries - 1:
                    wait_time = 5 * (attempt + 1)
                    print(f"  Waiting {wait_time} seconds before retry...")
                    time.sleep(wait_time)
                    continue
                return None

            # Check if transcript was truncated by comparing length
            original_words = len(transcript.split())
            formatted_words = len(result['transcript'].split())
            truncation_ratio = formatted_words / original_words

            if truncation_ratio < 0.9:  # If less than 90% of words are present
                print(f"  WARNING: Transcript appears truncated ({truncation_ratio*100:.1f}% of original)")
                print(f"  Using raw transcript with basic paragraph formatting instead")
                result['transcript'] = add_paragraph_breaks(transcript)

            return result

        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing failed on attempt {attempt + 1}/{max_retries}")
            logger.error(f"JSON error: {str(e)}")
            print(f"  ERROR: Failed to parse GPT-5 JSON response: {str(e)}")
            if attempt < max_retries - 1:
                wait_time = 5 * (attempt + 1)
                print(f"  Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                logger.error("GPT-5 formatting FAILED: JSON parse error after all retries")
                return None
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__

            # Log detailed error information
            logger.error(f"GPT-5 formatting attempt {attempt + 1}/{max_retries} failed")
            logger.error(f"Error type: {error_type}")
            logger.error(f"Error message: {error_msg}")
            logger.error(f"Transcript length: {len(transcript)} chars, {len(transcript.split())} words")

            # Check if it's a rate limit error
            if 'rate_limit' in error_msg.lower():
                wait_time = 60 * (attempt + 1)
                print(f"  Rate limit hit. Waiting {wait_time} seconds...")
                logger.warning(f"Rate limit on GPT-5, waiting {wait_time}s")
            elif '502' in error_msg or 'bad gateway' in error_msg.lower():
                wait_time = min(60 * (attempt + 1), 180)  # 60, 120, 180s for server errors
                print(f"  OpenAI server error (502 Bad Gateway). Waiting {wait_time} seconds...")
                logger.warning(f"502 Bad Gateway from OpenAI on GPT-5, waiting {wait_time}s")
            elif 'timeout' in error_msg.lower():
                wait_time = min(30 * (2 ** attempt), 120)
                print(f"  Timeout error. Waiting {wait_time} seconds...")
                logger.warning(f"Timeout on GPT-5, waiting {wait_time}s")
            elif 'connection' in error_msg.lower():
                wait_time = min(20 * (2 ** attempt), 120)
                print(f"  Connection error. Waiting {wait_time} seconds...")
                logger.warning(f"Connection error on GPT-5, waiting {wait_time}s")
            else:
                wait_time = min(10 * (2 ** attempt), 120)  # Exponential backoff, cap at 2 minutes
                print(f"  Error ({error_type}): {error_msg}")

            if attempt < max_retries - 1:
                print(f"  Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            else:
                print(f"  ERROR: GPT-5 formatting failed after {max_retries} attempts: {error_msg}")
                logger.error(f"GPT-5 formatting FAILED after all retries")
                return None

    return None


def create_word_document(formatted_data: Dict[str, str], output_path: str) -> bool:
    """
    Create a formatted Word document from the processed data.

    Args:
        formatted_data: Dictionary with 'title', 'summary', and 'transcript'
        output_path: Path to save the Word document

    Returns:
        True if successful, False otherwise
    """
    try:
        print(f"  Creating Word document...")

        doc = Document()

        # Title (Heading 1, Bold)
        title = doc.add_heading(formatted_data['title'], level=1)
        title.runs[0].bold = True

        # Add blank line
        doc.add_paragraph()

        # Summary heading (Heading 2)
        doc.add_heading('Summary', level=2)

        # Summary content
        doc.add_paragraph(formatted_data['summary'])

        # Add blank line
        doc.add_paragraph()

        # Transcript heading (Heading 2)
        doc.add_heading('Transcript', level=2)

        # Transcript content (split into paragraphs)
        transcript_text = formatted_data['transcript']

        # Handle different paragraph break formats
        if '\n\n' in transcript_text:
            # Double newlines - split on them
            paragraphs = transcript_text.split('\n\n')
        elif '\n' in transcript_text:
            # Single newlines - check if they look like paragraph breaks
            # (Usually after sentence endings with reasonable length)
            lines = transcript_text.split('\n')
            paragraphs = []
            current = []

            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line - paragraph break
                    if current:
                        paragraphs.append(' '.join(current))
                        current = []
                else:
                    current.append(line)

            if current:
                paragraphs.append(' '.join(current))
        else:
            # No newlines - add as single paragraph
            paragraphs = [transcript_text]

        # Add paragraphs to document
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

        # Save document
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"  ERROR: Failed to create Word document: {str(e)}")
        return False


def process_audio_file(input_path: Path, keep_compressed: bool = False) -> bool:
    """
    Process a single audio file through the complete pipeline.

    Args:
        input_path: Path to input audio file
        keep_compressed: Whether to keep compressed audio files

    Returns:
        True if successful, False otherwise
    """
    filename = input_path.stem
    output_path = Path(OUTPUT_DIR) / f"{filename}.docx"

    # Check if already processed
    if output_path.exists():
        print(f"\nSkipping: {input_path.name} (already processed)")
        return True

    print(f"\nProcessing: {input_path.name}")

    # Step 1: Compress audio
    compressed_path = Path(COMPRESSED_DIR) / f"{filename}.mp3"
    if not compress_audio(str(input_path), str(compressed_path)):
        return False

    # Step 2: Check file size
    is_valid, size_mb = check_file_size(str(compressed_path))
    print(f"  Compressed file size: {size_mb:.2f}MB")

    if not is_valid:
        print(f"  ERROR: File still too large after compression ({size_mb:.2f}MB > {MAX_FILE_SIZE_MB}MB)")
        if not keep_compressed:
            compressed_path.unlink()
        return False

    # Step 3: Transcribe
    transcript = transcribe_audio(str(compressed_path))
    if transcript is None:
        if not keep_compressed:
            compressed_path.unlink()
        return False

    # Step 4: Format with GPT-4
    formatted_data = format_with_gpt4(transcript)
    if formatted_data is None:
        if not keep_compressed:
            compressed_path.unlink()
        return False

    # Step 5: Create Word document
    output_path = Path(OUTPUT_DIR) / f"{filename}.docx"
    success = create_word_document(formatted_data, str(output_path))

    # Clean up compressed file if requested
    if not keep_compressed and compressed_path.exists():
        compressed_path.unlink()
        print(f"  Deleted compressed file")

    if success:
        print(f"  SUCCESS: Created {output_path.name}")

    return success


def main():
    """Main entry point for the application."""
    parser = argparse.ArgumentParser(
        description='Process audio files into formatted Word documents'
    )
    parser.add_argument(
        '--keep-compressed',
        action='store_true',
        help='Keep compressed audio files after processing'
    )
    parser.add_argument(
        '--file',
        type=str,
        help='Process a single specific file from the input directory'
    )
    args = parser.parse_args()

    print("=" * 60)
    print("Audio Transcription & Documentation Tool")
    print("=" * 60)

    # Log session start
    logger.info("="*60)
    logger.info("NEW PROCESSING SESSION STARTED")
    logger.info(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info("="*60)

    # Validate API key
    if not os.getenv('OPENAI_API_KEY'):
        print("\nERROR: OPENAI_API_KEY not found in environment variables")
        print("Please set it in a .env file or export it as an environment variable")
        logger.error("Missing OPENAI_API_KEY")
        return

    # Check network connectivity
    if not check_network_connectivity():
        print("\nWARNING: Network connectivity issues detected")
        print("Proceeding anyway, but expect connection errors...")
        logger.warning("Proceeding with network connectivity issues")

    # Ensure directories exist
    Path(INPUT_DIR).mkdir(exist_ok=True)
    Path(COMPRESSED_DIR).mkdir(exist_ok=True)
    Path(OUTPUT_DIR).mkdir(exist_ok=True)

    # Get all audio files from input directory
    input_path = Path(INPUT_DIR)
    audio_files = []
    for ext in SUPPORTED_FORMATS:
        audio_files.extend(input_path.glob(f"*{ext}"))

    # Filter to specific file if --file is provided
    if args.file:
        specific_file = Path(INPUT_DIR) / args.file
        if not specific_file.exists():
            # Try partial match
            matches = [f for f in audio_files if args.file.lower() in f.name.lower()]
            if len(matches) == 1:
                audio_files = matches
            elif len(matches) > 1:
                print(f"\nMultiple files match '{args.file}':")
                for m in matches:
                    print(f"  - {m.name}")
                return
            else:
                print(f"\nFile not found: {args.file}")
                return
        else:
            audio_files = [specific_file]

    if not audio_files:
        print(f"\nNo audio files found in '{INPUT_DIR}/' directory")
        print(f"Supported formats: {', '.join(SUPPORTED_FORMATS)}")
        return

    print(f"\nFound {len(audio_files)} audio file(s) to process")

    # Process each file
    successful = 0
    failed = 0
    skipped = 0

    for idx, audio_file in enumerate(audio_files):
        # Check if file was already processed before attempting
        output_path = Path(OUTPUT_DIR) / f"{audio_file.stem}.docx"
        was_already_done = output_path.exists()

        result = process_audio_file(audio_file, args.keep_compressed)

        if was_already_done:
            skipped += 1
        elif result:
            successful += 1
        else:
            failed += 1

        # Add delay between files to avoid rate limiting (but not after the last file)
        if idx < len(audio_files) - 1 and not was_already_done:
            wait_time = 15  # 15 seconds between files to reduce server load
            print(f"  Waiting {wait_time} seconds before next file...")
            time.sleep(wait_time)

    # Print summary
    print("\n" + "=" * 60)
    print("PROCESSING COMPLETE")
    print("=" * 60)
    print(f"Successful: {successful}")
    print(f"Failed: {failed}")
    print(f"Skipped (already processed): {skipped}")
    print(f"Total: {len(audio_files)}")

    # Log session summary
    logger.info("="*60)
    logger.info("PROCESSING SESSION COMPLETE")
    logger.info(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    logger.info(f"Results: {successful} successful, {failed} failed, {skipped} skipped out of {len(audio_files)} total")
    logger.info("="*60)

    if successful > 0:
        print(f"\nOutput documents saved to: {OUTPUT_DIR}/")

    if failed > 0:
        print(f"\nCheck transcription_errors.log for detailed error information")


if __name__ == "__main__":
    main()
